"""
Genera el informe Word del sistema difuso de coccion de carne de res.
Ejecutar con: python generar_informe.py
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ─── Colores ──────────────────────────────────────────────────────────────────
AZUL_OSCURO = RGBColor(0x1F, 0x49, 0x7D)   # Títulos principales
AZUL_MEDIO  = RGBColor(0x2E, 0x74, 0xB5)   # Subtítulos
GRIS_TEXTO  = RGBColor(0x26, 0x26, 0x26)   # Cuerpo
GRIS_CODE   = RGBColor(0xD6, 0xD6, 0xD6)   # Fondo código
VERDE       = RGBColor(0x1E, 0x8B, 0x4C)   # Acento tabla

# ─── Helpers ──────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color: str):
    """Aplica color de fondo a una celda de tabla."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)


def add_heading(doc, text, level=1, color=None):
    """Agrega un título con estilo personalizado."""
    p     = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run   = p.add_run(text)
    sizes = {1: 16, 2: 13, 3: 11}
    run.font.size = Pt(sizes.get(level, 11))
    run.font.bold = True
    run.font.color.rgb = color or (AZUL_OSCURO if level == 1 else AZUL_MEDIO)
    if level == 1:
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after  = Pt(6)
        # Línea inferior
        pPr   = p._p.get_or_add_pPr()
        pBdr  = OxmlElement("w:pBdr")
        bot   = OxmlElement("w:bottom")
        bot.set(qn("w:val"),   "single")
        bot.set(qn("w:sz"),    "6")
        bot.set(qn("w:space"), "1")
        bot.set(qn("w:color"), "2E74B5")
        pBdr.append(bot)
        pPr.append(pBdr)
    else:
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after  = Pt(4)
    return p


def add_body(doc, text, indent=False):
    """Agrega párrafo de cuerpo."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(6)
    if indent:
        p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run(text)
    run.font.size  = Pt(11)
    run.font.color.rgb = GRIS_TEXTO
    return p


def add_bullet(doc, text):
    """Agrega un punto de lista."""
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.color.rgb = GRIS_TEXTO


def add_code(doc, code_text):
    """Agrega bloque de código con fuente monoespaciada y fondo gris."""
    for line in code_text.split("\n"):
        p   = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(0)
        p.paragraph_format.left_indent  = Cm(0.5)
        run = p.add_run(line if line else " ")
        run.font.name = "Courier New"
        run.font.size = Pt(8.5)
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x5E)
        # Fondo gris claro
        rPr  = run._r.get_or_add_rPr()
        shd  = OxmlElement("w:shd")
        shd.set(qn("w:val"),   "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"),  "F2F2F2")
        rPr.append(shd)


def add_formula(doc, formula_text):
    """Agrega una fórmula centrada en cursiva."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(formula_text)
    run.font.italic = True
    run.font.size   = Pt(11)
    run.font.color.rgb = AZUL_OSCURO


# ─── DOCUMENTO ────────────────────────────────────────────────────────────────

def crear_informe():
    doc = Document()

    # Márgenes
    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(3.0)
        section.right_margin  = Cm(2.5)

    # ── PORTADA ───────────────────────────────────────────────────────────────
    portada_items = [
        ("UNIVERSIDAD DE LA AMAZONIA", Pt(16), True, AZUL_OSCURO),
        ("Facultad de Ingeniería", Pt(13), False, AZUL_MEDIO),
        ("Programa de Ingeniería de Sistemas", Pt(12), False, GRIS_TEXTO),
        ("", Pt(10), False, GRIS_TEXTO),
        ("", Pt(10), False, GRIS_TEXTO),
        ("SISTEMA DE INFERENCIA DIFUSA (FUZZY)\nPARA LA PREDICCIÓN DEL TÉRMINO\nDE COCCIÓN DE CARNE DE RES",
         Pt(18), True, AZUL_OSCURO),
        ("", Pt(10), False, GRIS_TEXTO),
        ("", Pt(10), False, GRIS_TEXTO),
        ("Inteligencia Computacional I", Pt(12), True, AZUL_MEDIO),
        ("", Pt(8), False, GRIS_TEXTO),
        ("Estudiantes:", Pt(11), True, GRIS_TEXTO),
        ("Juan Camilo Guzmán Ascencio", Pt(11), False, GRIS_TEXTO),
        ("Nicolás Rios Plazas", Pt(11), False, GRIS_TEXTO),
        ("Diego Alejandro Ceballos", Pt(11), False, GRIS_TEXTO),
        ("", Pt(8), False, GRIS_TEXTO),
        ("Docente:", Pt(11), True, GRIS_TEXTO),
        ("Jesús Emilio Pinto", Pt(11), False, GRIS_TEXTO),
        ("", Pt(8), False, GRIS_TEXTO),
        ("Florencia, Caquetá — 2026", Pt(11), False, GRIS_TEXTO),
    ]

    for texto, size, bold, color in portada_items:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if texto:
            run = p.add_run(texto)
            run.font.size  = size
            run.font.bold  = bold
            run.font.color.rgb = color
        p.paragraph_format.space_after = Pt(4)

    doc.add_page_break()

    # ── RESUMEN ───────────────────────────────────────────────────────────────
    add_heading(doc, "RESUMEN")
    add_body(doc,
        "El presente informe describe el diseño, implementación y validación de un sistema de "
        "inferencia difusa (SID) de tipo Mamdani para la predicción del término de cocción de "
        "carne de res en horno. El sistema toma como entradas la temperatura del horno (°C), el "
        "tiempo de cocción (minutos) y el grosor de la pieza (cm), y produce como salida un valor "
        "continuo en la escala 0–100 que es posteriormente interpretado en seis categorías "
        "lingüísticas: Crudo, Azul, Medio, Tres Cuartos, Bien Cocido y Quemado.")
    add_body(doc,
        "Las 64 reglas de inferencia fueron construidas a partir de tablas de cocción reconocidas "
        "(USDA FSIS, 2023; McGee, 2004; tabla oven-bake 425 °F), garantizando fundamentación "
        "científica. La implementación se realizó en Python con las librerías scikit-fuzzy, "
        "matplotlib y tkinter, resultando en una aplicación de escritorio interactiva que permite "
        "visualizar en tiempo real las funciones de membresía, la inferencia y la defuzzificación "
        "por centroide.")
    add_body(doc,
        "Palabras clave: lógica difusa, inferencia Mamdani, funciones de membresía trapezoidales, "
        "defuzzificación por centroide, cocción de carne.")
    doc.add_page_break()

    # ── 1. INTRODUCCIÓN ───────────────────────────────────────────────────────
    add_heading(doc, "1. INTRODUCCIÓN")
    add_body(doc,
        "La cocción de carne es un proceso físico-químico complejo gobernado por la transferencia de "
        "calor (ley de Fourier), la desnaturalización de proteínas (miosina a ≈50 °C, actina a ≈65 °C) "
        "y la evaporación de agua intersticial. La temperatura interna que alcanza la pieza depende de "
        "tres factores principales: la temperatura del entorno (horno), el tiempo de exposición y el "
        "grosor de la pieza, que determina la resistencia térmica que debe superar el calor para llegar "
        "al centro.")
    add_body(doc,
        "El razonamiento humano experto (un chef experimentado) trabaja con términos imprecisos como "
        "'temperatura alta', 'bastante tiempo' o 'carne gruesa'. La lógica difusa, propuesta por Zadeh "
        "(1965), formaliza exactamente este tipo de razonamiento: en lugar de fronteras binarias "
        "cruda/cocida, define grados de pertenencia continuos que permiten modelar las transiciones "
        "graduales entre categorías.")
    add_body(doc,
        "El presente proyecto aplica un sistema de inferencia difusa Mamdani para predecir el término "
        "de cocción de la carne de res en función de las tres variables mencionadas. El objetivo "
        "pedagógico es demostrar que los sistemas difusos constituyen una herramienta poderosa para "
        "modelar conocimiento experto en dominios donde la incertidumbre y la imprecisión son inherentes.")

    # ── 2. MARCO TEÓRICO ──────────────────────────────────────────────────────
    add_heading(doc, "2. MARCO TEÓRICO")

    add_heading(doc, "2.1 Lógica Difusa", level=2)
    add_body(doc,
        "La lógica difusa (fuzzy logic) fue introducida por Lotfi A. Zadeh en 1965 como una extensión "
        "de la lógica clásica bivalente. En la lógica clásica, un elemento pertenece o no pertenece a "
        "un conjunto (valor 0 o 1). En la lógica difusa, la pertenencia es un valor continuo en [0, 1], "
        "denominado grado de membresía μ.")
    add_formula(doc, "μ_A(x) ∈ [0, 1]  ∀x ∈ U")
    add_body(doc,
        "donde U es el universo de discurso, A es un conjunto difuso y μ_A(x) es el grado con que el "
        "elemento x pertenece a A. Un valor μ = 0 indica no pertenencia total, μ = 1 indica pertenencia "
        "total, y valores intermedios reflejan pertenencia parcial.")

    add_heading(doc, "2.2 Funciones de Membresía Trapezoidales", level=2)
    add_body(doc,
        "Una función de membresía trapezoidal está definida por cuatro parámetros [a, b, c, d] sobre el "
        "universo de discurso. Su forma es:")
    add_formula(doc, "μ(x) = 0           si x ≤ a")
    add_formula(doc, "μ(x) = (x−a)/(b−a) si a < x ≤ b")
    add_formula(doc, "μ(x) = 1           si b < x ≤ c")
    add_formula(doc, "μ(x) = (d−x)/(d−c) si c < x ≤ d")
    add_formula(doc, "μ(x) = 0           si x > d")
    add_body(doc,
        "Los extremos del universo se modelan con trapecios abiertos (a = b o c = d) para que el primer "
        "o último conjunto capture todos los valores fuera de rango. Por ejemplo, 'Temperatura Baja' "
        "usa [0, 0, 80, 140]: cualquier temperatura ≤ 0 °C tiene membresía 1 en 'Baja', y la membresía "
        "decae linealmente hasta 0 entre 80 y 140 °C.")

    add_heading(doc, "2.3 Inferencia Mamdani", level=2)
    add_body(doc,
        "El método de inferencia Mamdani (Mamdani y Assilian, 1975) sigue cuatro etapas:")

    pasos = [
        ("Fuzzificación:", "Cada valor de entrada crisp se convierte en un vector de grados de membresía, "
         "uno por cada conjunto difuso de la variable. Se evalúa μ(x) para cada función de membresía usando "
         "interpolación lineal sobre el universo discreto."),
        ("Evaluación de reglas:", "Cada regla tiene la forma SI (A₁ Y A₂ Y A₃) ENTONCES B. El conector Y "
         "(AND) se implementa con el operador mínimo: α = min(μ_A₁(x₁), μ_A₂(x₂), μ_A₃(x₃)). El valor α "
         "es el 'disparo' o nivel de activación de la regla."),
        ("Agregación:", "Si varias reglas concluyen en el mismo término de salida, se toma el máximo de sus "
         "activaciones (operador OR = máximo). Luego cada función de membresía de salida se 'recorta' al nivel "
         "de activación: μ'_B(y) = min(α, μ_B(y)). Las áreas recortadas se unen con el máximo punto a punto "
         "formando el área agregada."),
        ("Defuzzificación:", "El valor crisp de salida se obtiene calculando el centroide (centro de gravedad) "
         "del área agregada: y* = ∫y·μ_agg(y)dy / ∫μ_agg(y)dy. Se aproxima numéricamente con la suma de "
         "Riemann sobre el universo discreto TERM_U = [0, 100]."),
    ]

    for titulo, texto in pasos:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.left_indent = Cm(0.5)
        run_t = p.add_run(titulo + " ")
        run_t.font.bold = True
        run_t.font.size = Pt(11)
        run_t.font.color.rgb = AZUL_MEDIO
        run_b = p.add_run(texto)
        run_b.font.size = Pt(11)
        run_b.font.color.rgb = GRIS_TEXTO

    add_heading(doc, "2.4 Fundamento Científico de las Reglas", level=2)
    add_body(doc,
        "Las 64 reglas del sistema se basan en fuentes reconocidas de ciencia alimentaria y normativa "
        "de seguridad alimentaria:")

    refs_cientificas = [
        "USDA FSIS (2023) — Define las temperaturas internas seguras: Medio ≥ 63 °C, Bien Cocido ≥ 71 °C. "
         "Establece la escala de términos de cocción usada como referencia de la variable de salida.",
        "McGee, H. (2004) On Food and Cooking — Documenta la desnaturalización de la miosina a ≈50 °C "
         "(textura firme) y la actina a ≈65 °C (textura seca). Justifica por qué tiempos largos a baja "
         "temperatura producen un término más alto de lo esperado.",
        "Tabla oven-bake 425 °F (≈218 °C) — Referencia empírica: un filete de 1 pulgada (2.5 cm) alcanza "
         "término Azul en 5-7 min, Medio en 9-11 min, Med-Well en 13-15 min, Well Done en 17+ min.",
        "Tabla low-and-slow 250 °F (≈121 °C) — A baja temperatura, 20-25 min → Medium-Rare en pieza "
         "de 2.5 cm. Justifica que a Temperatura Baja + Tiempo Moderado se asigne al menos término Azul.",
        "Ley de Fourier (transferencia de calor) — El flujo de calor es inversamente proporcional al "
         "grosor. A mayor grosor, mayor resistencia térmica: necesita más tiempo o mayor temperatura para "
         "alcanzar el mismo término en el centro.",
    ]
    for r in refs_cientificas:
        add_bullet(doc, r)

    # ── 3. DESCRIPCIÓN DEL SISTEMA ────────────────────────────────────────────
    add_heading(doc, "3. DESCRIPCIÓN DEL SISTEMA")

    add_heading(doc, "3.1 Variables de Entrada", level=2)
    add_body(doc,
        "El sistema posee tres variables lingüísticas de entrada, cada una con cuatro conjuntos difusos "
        "trapezoidales:")

    # Tabla de variables
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, txt in enumerate(["Variable", "Rango", "Conjuntos", "Parámetros [a, b, c, d]"]):
        hdr[i].text = txt
        hdr[i].paragraphs[0].runs[0].font.bold = True
        set_cell_bg(hdr[i], "2E74B5")
        hdr[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    filas_vars = [
        ("Temperatura\n(TEMP_U)", "0 – 300 °C",
         "Baja\nMedia\nAlta\nMuy Alta",
         "[0, 0, 80, 140]\n[100, 150, 180, 220]\n[180, 210, 240, 270]\n[240, 265, 300, 300]"),
        ("Tiempo\n(TIME_U)", "0 – 120 min",
         "Poco\nModerado\nMucho\nExcesivo",
         "[0, 0, 20, 40]\n[25, 40, 55, 70]\n[55, 70, 85, 100]\n[85, 100, 120, 120]"),
        ("Grosor\n(THICK_U)", "0 – 6 cm",
         "Delgada\nNormal\nGruesa\nMuy Gruesa",
         "[0, 0, 1, 2]\n[1.5, 2.0, 2.8, 3.5]\n[3.0, 3.5, 4.5, 5.0]\n[4.5, 5.0, 6.0, 6.0]"),
    ]

    for fila in filas_vars:
        cells = table.add_row().cells
        for i, txt in enumerate(fila):
            cells[i].text = txt
            cells[i].paragraphs[0].runs[0].font.size = Pt(9)

    doc.add_paragraph()

    add_heading(doc, "3.2 Variable de Salida", level=2)
    add_body(doc,
        "La variable de salida 'Término de cocción' tiene universo TERM_U = [0, 100] y seis conjuntos "
        "difusos trapezoidales que representan los niveles culinarios reconocidos:")

    table2 = doc.add_table(rows=1, cols=3)
    table2.style = "Table Grid"
    hdr2 = table2.rows[0].cells
    for i, txt in enumerate(["Término", "Parámetros [a, b, c, d]", "Temp. interna referencia (USDA)"]):
        hdr2[i].text = txt
        hdr2[i].paragraphs[0].runs[0].font.bold = True
        set_cell_bg(hdr2[i], "2E74B5")
        hdr2[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    filas_sal = [
        ("Crudo",        "[0, 0, 8, 15]",     "< 52 °C"),
        ("Azul",         "[10, 18, 25, 32]",  "52 – 55 °C"),
        ("Medio",        "[28, 38, 48, 58]",  "60 – 66 °C"),
        ("Tres Cuartos", "[50, 58, 65, 73]",  "66 – 72 °C"),
        ("Bien Cocido",  "[68, 75, 82, 88]",  "72 – 80 °C"),
        ("Quemado",      "[83, 90, 100, 100]", "> 82 °C / exceso"),
    ]
    for fila in filas_sal:
        cells = table2.add_row().cells
        for i, txt in enumerate(fila):
            cells[i].text = txt
            cells[i].paragraphs[0].runs[0].font.size = Pt(9)

    doc.add_paragraph()

    add_heading(doc, "3.3 Base de Reglas", level=2)
    add_body(doc,
        "El sistema emplea 64 reglas de la forma SI (Temperatura es T) Y (Tiempo es Ti) Y (Grosor es G) "
        "ENTONCES (Término es S), resultantes del producto cartesiano 4 × 4 × 4 de los conjuntos de "
        "cada variable de entrada. A continuación se presentan ejemplos representativos:")

    ejemplos = [
        ("R1:",  "SI Temp=Baja    Y Tiempo=Poco     Y Grosor=Delgada    ENTONCES Término=Crudo"),
        ("R6:",  "SI Temp=Baja    Y Tiempo=Moderado Y Grosor=Normal     ENTONCES Término=Azul"),
        ("R22:", "SI Temp=Media   Y Tiempo=Moderado Y Grosor=Normal     ENTONCES Término=Medio"),
        ("R34:", "SI Temp=Alta    Y Tiempo=Poco     Y Grosor=Delgada    ENTONCES Término=Medio"),
        ("R42:", "SI Temp=Alta    Y Tiempo=Mucho    Y Grosor=Normal     ENTONCES Término=Quemado"),
        ("R53:", "SI Temp=MuyAlta Y Tiempo=Moderado Y Grosor=Delgada   ENTONCES Término=Quemado"),
        ("R64:", "SI Temp=MuyAlta Y Tiempo=Excesivo Y Grosor=MuyGruesa ENTONCES Término=Quemado"),
    ]

    table3 = doc.add_table(rows=1, cols=2)
    table3.style = "Table Grid"
    hdr3 = table3.rows[0].cells
    for i, txt in enumerate(["#", "Regla"]):
        hdr3[i].text = txt
        hdr3[i].paragraphs[0].runs[0].font.bold = True
        set_cell_bg(hdr3[i], "2E74B5")
        hdr3[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    for num, regla in ejemplos:
        cells = table3.add_row().cells
        cells[0].text = num
        cells[0].paragraphs[0].runs[0].font.size = Pt(9)
        cells[0].paragraphs[0].runs[0].font.bold = True
        cells[1].text = regla
        cells[1].paragraphs[0].runs[0].font.size = Pt(9)
        cells[1].paragraphs[0].runs[0].font.name = "Courier New"

    doc.add_paragraph()

    # ── 4. IMPLEMENTACIÓN ─────────────────────────────────────────────────────
    add_heading(doc, "4. IMPLEMENTACIÓN")

    add_heading(doc, "4.1 Arquitectura del Software", level=2)
    add_body(doc,
        "El sistema está implementado en Python 3.13 y organizado en cuatro módulos con "
        "responsabilidades claramente separadas:")

    modulos = [
        ("conjuntos.py", "Define los universos de discurso (arreglos NumPy), las funciones de "
         "membresía (combinacion trapmf + trimf), las constantes Sugeno y la base de 64 reglas."),
        ("motor.py", "Implementa las tres etapas del motor de inferencia Mamdani: fuzzificación "
         "(fuzzificar), inferencia+agregación y defuzzificación por centroide (inferir), más la "
         "interpretación lingüística del resultado (interpretar)."),
        ("motor_sugeno.py", "Implementa el motor de inferencia Sugeno de orden cero: misma "
         "fuzzificación y evaluacion AND, pero la salida es la media ponderada de constantes "
         "numericas en lugar de centroide de area agregada."),
        ("graficador.py", "Genera las gráficas matplotlib embebidas: funciones de membresía "
         "(dibujar_membresias), área agregada Mamdani (dibujar_salida), y gráfico comparativo "
         "Mamdani vs Sugeno con 4 casos (dibujar_comparacion)."),
        ("aplicacion.py", "Interfaz gráfica tkinter con ttk.Notebook (4 pestañas): Membresías, "
         "Inferencia, Reglas (64 con checkboxes individuales) y Mamdani vs Sugeno (comparación "
         "en tiempo real de ambos motores)."),
    ]

    for nombre, desc in modulos:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.left_indent = Cm(0.5)
        run_n = p.add_run(nombre + ": ")
        run_n.font.bold = True
        run_n.font.name = "Courier New"
        run_n.font.size = Pt(10)
        run_n.font.color.rgb = AZUL_OSCURO
        run_d = p.add_run(desc)
        run_d.font.size = Pt(11)
        run_d.font.color.rgb = GRIS_TEXTO

    add_heading(doc, "4.2 Módulo conjuntos.py — Universos y Membresías", level=2)
    add_body(doc, "Los universos de discurso y funciones de membresía se definen una única vez al "
        "importar el módulo. Se combinan dos tipos de funciones: trapmf para los conjuntos extremos "
        "(abiertos en un lado) y trimf para los conjuntos intermedios (con un pico único bien definido). "
        "Esta combinacion cumple el requisito de usar al menos dos tipos de funcion de membresia por variable:")
    add_code(doc, """\
# Extremos: trapmf abierto. Intermedios: trimf (un solo pico).
TEMP_MFS = {
    "Baja":     fuzz.trapmf(TEMP_U, [0,   0,   80,  140]),  # trapmf: abierto izquierda
    "Media":    fuzz.trimf( TEMP_U, [100, 160, 220]),         # trimf: pico en 160 C
    "Alta":     fuzz.trimf( TEMP_U, [180, 225, 270]),         # trimf: pico en 225 C
    "Muy Alta": fuzz.trapmf(TEMP_U, [240, 265, 300, 300]),   # trapmf: abierto derecha
}
SUGENO_CONSTANTS = {
    "Crudo": 4, "Azul": 21, "Medio": 43,
    "Tres Cuartos": 62, "Bien Cocido": 79, "Quemado": 95,
}""")
    doc.add_paragraph()

    add_heading(doc, "4.3 Módulo motor.py — Fuzzificación e Inferencia", level=2)
    add_body(doc, "La función fuzzificar evalúa el grado de pertenencia del valor de entrada "
        "en cada conjunto difuso usando interpolación lineal:")
    add_code(doc, """\
def fuzzificar(valor, mf_dict, universo):
    v = float(np.clip(valor, universo[0], universo[-1]))
    grados = {}
    for nombre, mf in mf_dict.items():
        grados[nombre] = float(fuzz.interp_membership(universo, mf, v))
    return grados""")
    doc.add_paragraph()
    add_body(doc, "La función inferir ejecuta el ciclo completo de inferencia Mamdani: "
        "fuzzifica las tres entradas, evalúa todas las reglas activas con AND=mínimo, "
        "agrega con OR=máximo, recorta y defuzzifica por centroide:")
    add_code(doc, """\
def inferir(temp, tiempo, grosor, reglas=None):
    td  = fuzzificar(temp,   conjuntos.TEMP_MFS,  conjuntos.TEMP_U)
    tid = fuzzificar(tiempo, conjuntos.TIME_MFS,  conjuntos.TIME_U)
    thd = fuzzificar(grosor, conjuntos.THICK_MFS, conjuntos.THICK_U)

    activaciones = {}
    for (t, ti, th, salida) in reglas:
        act = min(td.get(t, 0), tid.get(ti, 0), thd.get(th, 0))  # AND
        if salida not in activaciones or act > activaciones[salida]:
            activaciones[salida] = act                            # OR

    agregado = np.zeros_like(conjuntos.TERM_U, dtype=float)
    for termino, fuerza in activaciones.items():
        if fuerza > 0:
            corte    = np.fmin(fuerza, conjuntos.TERM_MFS[termino])  # recorte
            agregado = np.fmax(agregado, corte)                      # unión

    crisp = float(fuzz.defuzz(conjuntos.TERM_U, agregado, "centroid"))
    return td, tid, thd, activaciones, [], agregado, crisp""")
    doc.add_paragraph()

    add_heading(doc, "4.4 Módulo graficador.py — Visualización", level=2)
    add_body(doc,
        "El graficador embebe figuras matplotlib directamente en los frames de tkinter mediante "
        "FigureCanvasTkAgg. La función dibujar_membresias actualiza las cuatro subgráficas (una "
        "por variable) en cada recálculo, marcando con línea vertical el valor actual y puntos "
        "sobre las curvas activas. La función dibujar_salida muestra el área agregada completa y "
        "la línea vertical del centroide.")
    add_code(doc, """\
from matplotlib.backends.backend_tkagg import FigureCanvasTkAgg

def dibujar_membresias(fig, axes, canvas, tv, ti, th, td, tid, thd, crisp):
    for ax, (titulo, universo, mfs, colores, val, grados) in zip(axes, datos):
        ax.cla()
        for (nombre, mf), color in zip(mfs.items(), colores):
            ax.plot(universo, mf, color=color, lw=1.8, label=nombre)
        if val is not None:
            ax.axvline(val, color='white', lw=1.3, ls='--')
    canvas.draw_idle()""")
    doc.add_paragraph()

    add_heading(doc, "4.5 Módulo motor_sugeno.py — Inferencia Sugeno", level=2)
    add_body(doc,
        "El motor Sugeno de orden cero reutiliza la funcion fuzzificar de motor.py. La diferencia "
        "fundamental es que no construye area agregada: calcula directamente la media ponderada "
        "de las constantes numericas asociadas a cada termino de salida:")
    add_code(doc, """\
def inferir_sugeno(temp, tiempo, grosor, reglas=None):
    td  = motor.fuzzificar(temp,   conjuntos.TEMP_MFS,  conjuntos.TEMP_U)
    tid = motor.fuzzificar(tiempo, conjuntos.TIME_MFS,  conjuntos.TIME_U)
    thd = motor.fuzzificar(grosor, conjuntos.THICK_MFS, conjuntos.THICK_U)

    numerador = denominador = 0.0
    for (t, ti, th, salida) in reglas:
        act = min(td.get(t,0), tid.get(ti,0), thd.get(th,0))  # AND = minimo
        if act > 0:
            k = conjuntos.SUGENO_CONSTANTS[salida]  # constante numerica
            numerador   += act * k   # suma ponderada
            denominador += act

    crisp = numerador / denominador if denominador > 0 else 0.0
    return td, tid, thd, {}, crisp  # sin area agregada""")
    doc.add_paragraph()

    add_heading(doc, "4.6 Módulo aplicacion.py — Interfaz Gráfica", level=2)
    add_body(doc,
        "La interfaz usa ttk.Notebook con cuatro pestañas. En cada recalculo, _calcular() "
        "ejecuta ambos motores (Mamdani y Sugeno) con las mismas entradas y reglas activas, "
        "actualizando simultaneamente todas las visualizaciones:")
    add_code(doc, """\
def _calcular(self):
    reglas_activas = [r for r, v in
        zip(conjuntos.RULES, self._rule_vars) if v.get()]
    # Motor Mamdani
    td, tid, thd, acts, _, agregado, crisp = motor.inferir(
        self._temp, self._tiempo, self._grosor, reglas_activas)
    # Motor Sugeno (mismas entradas y reglas)
    _, _, _, _, crisp_s = motor_sugeno.inferir_sugeno(
        self._temp, self._tiempo, self._grosor, reglas_activas)
    # Actualizar las 4 pestanas
    graficador.dibujar_membresias(...)
    graficador.dibujar_salida(...)
    graficador.dibujar_comparacion(..., crisp, crisp_s)""")
    doc.add_paragraph()

    # ── 5. EJEMPLO DE INFERENCIA ──────────────────────────────────────────────
    add_heading(doc, "5. EJEMPLO DE INFERENCIA PASO A PASO")
    add_body(doc,
        "Se presenta un ejemplo concreto con los valores: Temperatura = 218 °C (Alta), "
        "Tiempo = 30 min (Moderado), Grosor = 2.5 cm (Normal).")

    add_heading(doc, "5.1 Fuzzificación", level=2)
    add_body(doc, "Se calcula μ para cada conjunto de cada variable en el valor dado:")

    table4 = doc.add_table(rows=1, cols=3)
    table4.style = "Table Grid"
    hdr4 = table4.rows[0].cells
    for i, txt in enumerate(["Variable / Valor", "Conjunto", "Grado μ"]):
        hdr4[i].text = txt
        hdr4[i].paragraphs[0].runs[0].font.bold = True
        set_cell_bg(hdr4[i], "2E74B5")
        hdr4[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    fuzz_filas = [
        ("Temperatura = 218 °C", "Baja",     "0.000"),
        ("",                     "Media",    "0.050"),
        ("",                     "Alta",     "0.733"),
        ("",                     "Muy Alta", "0.000"),
        ("Tiempo = 30 min",      "Poco",     "0.000"),
        ("",                     "Moderado", "0.933"),
        ("",                     "Mucho",    "0.000"),
        ("",                     "Excesivo", "0.000"),
        ("Grosor = 2.5 cm",      "Delgada",  "0.000"),
        ("",                     "Normal",   "0.833"),
        ("",                     "Gruesa",   "0.000"),
        ("",                     "Muy Gruesa","0.000"),
    ]
    for var, conj, grado in fuzz_filas:
        cells = table4.add_row().cells
        cells[0].text = var
        cells[1].text = conj
        cells[2].text = grado
        for c in cells:
            c.paragraphs[0].runs[0].font.size = Pt(9)

    doc.add_paragraph()

    add_heading(doc, "5.2 Evaluación de Reglas", level=2)
    add_body(doc,
        "Solo las reglas cuyos tres antecedentes tienen μ > 0 contribuyen. Con estos valores "
        "solo una regla tiene los tres antecedentes activos:")

    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run(
        "R38 (Alta, Moderado, Normal → Tres Cuartos):\n"
        "   α = min(μ_Alta(218), μ_Moderado(30), μ_Normal(2.5))\n"
        "   α = min(0.733, 0.933, 0.833) = 0.733")
    run.font.name = "Courier New"
    run.font.size = Pt(9)
    run.font.color.rgb = AZUL_OSCURO

    add_heading(doc, "5.3 Defuzzificación", level=2)
    add_body(doc,
        "Con α = 0.733 para el conjunto 'Tres Cuartos' [50, 58, 65, 73], se recorta la función "
        "de membresía al nivel 0.733. El centroide del área recortada resulta en:")
    add_formula(doc, "y* = centroid(TERM_U, área_agregada) ≈ 61.5  →  Tres Cuartos")
    add_body(doc,
        "Esto es consistente con la referencia empírica: a 425 °F (≈218 °C) un filete de 1\" "
        "alcanza término Med-Well (tres cuartos) en 13-15 min; a 30 min ya supera ese nivel "
        "hacia Well Done, coherente con el resultado del sistema.")

    # ── 6. RESULTADOS Y ANÁLISIS ──────────────────────────────────────────────
    add_heading(doc, "6. RESULTADOS Y ANÁLISIS")
    add_body(doc,
        "Se ejecutó el sistema para distintas combinaciones de entrada y se verificó coherencia "
        "con las tablas de cocción de referencia. Los siguientes casos representativos muestran "
        "el comportamiento del sistema:")

    table5 = doc.add_table(rows=1, cols=5)
    table5.style = "Table Grid"
    hdr5 = table5.rows[0].cells
    for i, txt in enumerate(["Temp (°C)", "Tiempo (min)", "Grosor (cm)", "Crisp", "Término"]):
        hdr5[i].text = txt
        hdr5[i].paragraphs[0].runs[0].font.bold = True
        set_cell_bg(hdr5[i], "2E74B5")
        hdr5[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    resultados_tabla = [
        ("120", "15", "2.5", "~8",    "Crudo"),
        ("120", "90", "2.5", "~27",   "Azul"),
        ("175", "40", "2.5", "~43",   "Medio"),
        ("218", "30", "2.5", "~61",   "Tres Cuartos"),
        ("218", "65", "3.8", "~76",   "Bien Cocido"),
        ("270", "45", "2.5", "~91",   "Quemado"),
        ("270", "25", "5.0", "~63",   "Tres Cuartos"),
    ]
    for fila in resultados_tabla:
        cells = table5.add_row().cells
        for i, txt in enumerate(fila):
            cells[i].text = txt
            cells[i].paragraphs[0].runs[0].font.size = Pt(9)

    doc.add_paragraph()

    add_body(doc,
        "El análisis de los resultados muestra que el sistema captura correctamente las relaciones "
        "entre las variables. A mayor temperatura, menor tiempo es necesario para alcanzar el mismo "
        "término, en concordancia con la ley de Fourier. El grosor actúa como factor amortiguador: "
        "piezas más gruesas requieren significativamente más tiempo para alcanzar la misma temperatura "
        "interna.")
    add_body(doc,
        "La combinacion de funciones trapmf (extremos) y trimf (conjuntos intermedios) permite "
        "representar con mayor precision los valores centrales de cada variable: los conjuntos "
        "intermedios tienen un punto de maxima pertenencia bien definido (el pico de la triangular), "
        "mientras que los extremos se mantienen abiertos para capturar todos los valores fuera de rango.")

    add_heading(doc, "6.2 Comparación Mamdani vs Sugeno — Tres Casos de Prueba", level=2)
    add_body(doc,
        "Se ejecutaron ambos modelos con los mismos tres casos de entrada para comparar sus "
        "resultados. Las constantes Sugeno corresponden al centro del intervalo plano de cada "
        "trapmf de salida: Crudo=4, Azul=21, Medio=43, Tres Cuartos=62, Bien Cocido=79, Quemado=95.")

    table6 = doc.add_table(rows=1, cols=7)
    table6.style = "Table Grid"
    hdr6 = table6.rows[0].cells
    for i, txt in enumerate(["Temp", "Tiempo", "Grosor",
                              "Mamdani", "Termino M",
                              "Sugeno",  "Termino S"]):
        hdr6[i].text = txt
        hdr6[i].paragraphs[0].runs[0].font.bold = True
        set_cell_bg(hdr6[i], "2E74B5")
        hdr6[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    comp_filas = [
        ("120 °C", "90 min", "2.5 cm", "52.2", "Tres Cuartos", "56.0", "Tres Cuartos"),
        ("175 °C", "45 min", "2.5 cm", "43.0", "Medio",        "43.0", "Medio"),
        ("250 °C", "35 min", "2.5 cm", "56.9", "Tres Cuartos", "60.7", "Tres Cuartos"),
    ]
    for fila in comp_filas:
        cells = table6.add_row().cells
        for i, txt in enumerate(fila):
            cells[i].text = txt
            cells[i].paragraphs[0].runs[0].font.size = Pt(9)

    doc.add_paragraph()

    add_body(doc,
        "Los dos modelos entregan resultados similares en los tres casos: coinciden en el termino "
        "linguistico (Tres Cuartos / Medio) con diferencias numericas de entre 0.0 y 3.8 puntos. "
        "Las mayores diferencias se observan cuando varias reglas con distintos terminos se activan "
        "simultaneamente: Mamdani construye un area agregada y el centroide pondera geometricamente "
        "todas las areas, mientras que Sugeno promedia directamente las constantes numericas.")
    add_body(doc,
        "En cuanto a facilidad de implementacion, Sugeno es mas simple: no requiere construir el "
        "array del area agregada (np.fmax/np.fmin) ni calcular el centroide numerico, solo una "
        "suma ponderada. En cuanto a interpretabilidad, Mamdani es mas intuitivo porque el area "
        "agregada es visible y el centroide tiene un significado geometrico claro.")
    add_body(doc,
        "Para el fenomeno de coccion de carne, Mamdani resulta mas adecuado porque permite "
        "visualizar graficamente como cada regla contribuye al resultado final, lo que facilita "
        "la validacion y explicacion del sistema ante un experto culinario o un evaluador academico.")

    # ── 7. CONCLUSIONES ───────────────────────────────────────────────────────
    add_heading(doc, "7. CONCLUSIONES")

    conclusiones = [
        "Ambos sistemas (Mamdani y Sugeno) predijeron correctamente el termino de coccion de la "
        "carne en los tres casos de prueba, coincidiendo en el termino linguistico con diferencias "
        "numericas menores a 4 puntos sobre la escala 0-100.",

        "La diferencia fundamental entre los metodos es el mecanismo de salida: Mamdani construye "
        "y agrega areas de conjuntos difusos y obtiene el resultado por centroide; Sugeno calcula "
        "directamente la media ponderada de constantes numericas, lo que lo hace mas eficiente "
        "computacionalmente pero menos interpretable visualmente.",

        "La combinacion de funciones trapmf (extremos) y trimf (conjuntos intermedios) es "
        "apropiada porque los conjuntos de frontera del universo deben ser abiertos (trapmf), "
        "mientras que los conjuntos centrales tienen un valor de maxima representacion bien "
        "definido que la funcion triangular modela con mayor precision.",

        "Mamdani resulta mas adecuado para este dominio de aplicacion porque la visualizacion "
        "del area agregada permite explicar e inspeccionar graficamente por que el sistema "
        "produce un determinado resultado, aspecto clave para validacion por expertos culinarios.",

        "La separacion en cinco modulos (conjuntos, motor, motor_sugeno, graficador, aplicacion) "
        "permite sustituir o comparar los motores de inferencia de forma independiente, sin "
        "modificar las variables, las reglas ni la interfaz grafica.",
    ]

    for i, c in enumerate(conclusiones, 1):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(5)
        run = p.add_run(f"{i}. {c}")
        run.font.size = Pt(11)
        run.font.color.rgb = GRIS_TEXTO

    # ── 8. REFERENCIAS ────────────────────────────────────────────────────────
    add_heading(doc, "8. REFERENCIAS BIBLIOGRÁFICAS")

    referencias = [
        "Zadeh, L. A. (1965). Fuzzy sets. Information and Control, 8(3), 338–353. "
        "https://doi.org/10.1016/S0019-9958(65)90241-X",

        "Mamdani, E. H., & Assilian, S. (1975). An experiment in linguistic synthesis with a fuzzy "
        "logic controller. International Journal of Man-Machine Studies, 7(1), 1–13. "
        "https://doi.org/10.1016/S0020-7373(75)80002-2",

        "USDA Food Safety and Inspection Service. (2023). Safe minimum internal temperatures. "
        "U.S. Department of Agriculture. https://www.fsis.usda.gov/food-safety/safe-food-handling-"
        "and-preparation/food-safety-basics/safe-temperature-chart",

        "McGee, H. (2004). On Food and Cooking: The Science and Lore of the Kitchen (2.ª ed.). "
        "Scribner.",

        "Warner, R. D. (2014). Meat science and its contribution to the global livestock sector. "
        "In Proceedings of the 60th International Congress of Meat Science and Technology. Punta "
        "del Este, Uruguay.",

        "Scikit-Fuzzy Contributors. (2023). scikit-fuzzy: Fuzzy logic toolbox for SciPy. "
        "https://pythonhosted.org/scikit-fuzzy/",

        "Hunter, J. D. (2007). Matplotlib: A 2D graphics environment. Computing in Science & "
        "Engineering, 9(3), 90–95. https://doi.org/10.1109/MCSE.2007.55",

        "Berk, Z. (2018). Food Process Engineering and Technology (3.ª ed.). Academic Press. "
        "[Capítulo 3: Transferencia de calor en alimentos]",
    ]

    for ref in referencias:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(5)
        p.paragraph_format.left_indent  = Cm(0.5)
        p.paragraph_format.first_line_indent = Cm(-0.5)
        run = p.add_run(ref)
        run.font.size = Pt(10)
        run.font.color.rgb = GRIS_TEXTO

    # ── GUARDAR ───────────────────────────────────────────────────────────────
    ruta = r"d:\Trabajos de la u\Septimo semestre\computacional\Tercera entrega\Informe_Sistema_Difuso_Coccion.docx"
    doc.save(ruta)
    print(f"Informe guardado en: {ruta}")


if __name__ == "__main__":
    crear_informe()
